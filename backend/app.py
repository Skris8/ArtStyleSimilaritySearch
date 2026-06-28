"""API de ArtSim: catálogo visual, similaridad, reportes y administración."""

import base64
from collections import Counter
from datetime import datetime
import hashlib
import io
import os
from pathlib import Path
import re
import typing
import unicodedata

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import mysql.connector
from mysql.connector import Error
from PIL import Image
from qdrant_client import QdrantClient
import toml


# Rutas y parámetros compartidos por el catálogo, los reportes y el frontend.
ROOT = Path(__file__).resolve().parents[1]
DATASET_DIR = ROOT / 'ipynb' / 'ArtStylesDataset'
FRONTEND_DIR = ROOT / 'frontend'
QDRANT_COLLECTION = "ArtStyles_images"
REPORT_TEMPLATE_PATH = ROOT / 'backend' / 'templates' / 'Plantilla_Consulta_Similaridad_ArtStyleSimilaritySearch.docx'
REPORTS_DIR = ROOT / 'exports' / 'reports'
DOCX_MEDIA_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
MPLCONFIG_DIR = ROOT / '.cache' / 'matplotlib'
STYLE_CHART_COLORS = {
    'anime': '#F7A1F9',
    'caricature': '#F28C28',
    'cartoon': '#6CE2FF',
    'cubism': '#B08D57',
    'cyberpunk': '#FF6EC7',
    'digital_potrait': '#C97B84',
    'digital_portrait': '#C97B84',
    'digital_scenery': '#82C8FF',
    'sketch': '#5D5B56',
    'surrealism': '#4F6DCE',
}
STYLE_CHART_FALLBACK_COLORS = ('#76B7B2', '#EDC948', '#B07AA1', '#59A14F', '#9C755F')

try:
    MPLCONFIG_DIR.mkdir(parents=True, exist_ok=True)
    os.environ.setdefault('MPLCONFIGDIR', str(MPLCONFIG_DIR))
except Exception:
    # Matplotlib puede continuar con su ubicación de caché predeterminada.
    pass

load_dotenv(dotenv_path=ROOT / 'ipynb' / '.env')

app = FastAPI(title="ArtSim Backend")

# El frontend se publica desde el mismo proceso para evitar otro servidor web.
if FRONTEND_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="frontend_static")

@app.get("/")
def read_index():
    """Sirve la vista principal del catálogo."""

    index_file = FRONTEND_DIR / 'index.html'
    if not index_file.exists():
        raise HTTPException(status_code=404, detail="Index file not found")
    return FileResponse(index_file)


@app.get("/admin")
def read_admin():
    """Sirve el formulario de acceso administrativo."""

    admin_file = FRONTEND_DIR / 'admin.html'
    if not admin_file.exists():
        raise HTTPException(status_code=404, detail="Admin file not found")
    return FileResponse(admin_file)


@app.get("/about")
def read_about():
    """Sirve la página informativa y su formulario de contacto."""

    about_file = FRONTEND_DIR / 'about.html'
    if not about_file.exists():
        raise HTTPException(status_code=404, detail="About file not found")
    return FileResponse(about_file)


def image_to_base64(p: Path, max_size=(600,400)):
    """Devuelve una copia JPEG reducida y su resolución para la API local."""

    img = Image.open(p).convert("RGB")
    img.thumbnail(max_size)
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=85)
    data = base64.b64encode(buf.getvalue()).decode('ascii')
    return f"data:image/jpeg;base64,{data}", img.size


def image_to_bytes(p: Path, max_size=(900,700)):
    """Carga una imagen local en memoria y limita su tamaño para un reporte."""

    img = Image.open(p).convert("RGB")
    img.thumbnail(max_size)
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=88)
    buf.seek(0)
    return buf

# Conversión de imágenes y normalización de datos recibidos desde Qdrant.
def _normalize_search_text(value):
    """Normaliza mayúsculas y acentos para realizar búsquedas tolerantes."""

    if not value:
        return ''
    if not isinstance(value, str):
        value = str(value)
    normalized = unicodedata.normalize('NFKD', value)
    normalized = normalized.encode('ascii', 'ignore').decode('ascii')
    return normalized.lower().strip()


def normalize_base64_payload(value):
    """Asegura que una cadena base64 incluya el prefijo de URL de datos."""

    if not value:
        return None
    if isinstance(value, str) and value.startswith('data:'):
        return value
    if isinstance(value, str):
        return f"data:image/jpeg;base64,{value}"
    return None


def base64_to_image_stream(value):
    """Decodifica una imagen del payload y la prepara como JPEG en memoria."""

    if not value:
        return None
    if isinstance(value, str) and ',' in value:
        value = value.split(',', 1)[1]
    try:
        raw = base64.b64decode(value)
        img = Image.open(io.BytesIO(raw)).convert("RGB")
        img.thumbnail((900, 700))
        out = io.BytesIO()
        img.save(out, format='JPEG', quality=88)
        out.seek(0)
        return out
    except Exception:
        return None


def _safe_filename_part(value):
    """Convierte un identificador en un fragmento seguro para nombres de archivo."""

    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", str(value or "record")).strip("_")
    return cleaned[:48] or "record"


def _coerce_qdrant_id(value):
    """Convierte identificadores numéricos sin alterar UUID u otras claves."""

    try:
        if isinstance(value, str) and value.isdigit():
            return int(value)
    except Exception:
        pass
    return value


def _ids_match(left, right):
    """Compara IDs sin depender de si el cliente los serializó como texto."""

    return str(left) == str(right)


def _point_attr(point, name, default=None):
    """Lee de forma uniforme puntos representados por objetos o diccionarios."""

    if isinstance(point, dict):
        return point.get(name, default)
    return getattr(point, name, default)


def _iter_qdrant_points(result):
    """Extrae puntos de las formas de respuesta usadas por qdrant-client."""

    if result is None:
        return []
    if isinstance(result, tuple):
        result = result[0]
    if isinstance(result, dict):
        points = result.get('points')
        if points is not None:
            return points
        nested = result.get('result')
        if isinstance(nested, dict):
            return nested.get('points') or nested.get('result') or []
        return nested or []
    for attr in ('points', 'result'):
        points = getattr(result, attr, None)
        if points is not None:
            if isinstance(points, dict):
                return points.get('points') or points.get('result') or []
            return points
    try:
        return list(result)
    except TypeError:
        return []


def _score_percent(score):
    """Convierte una puntuación normalizada a un porcentaje acotado."""

    if score is None:
        return None
    try:
        return round(max(0.0, min(float(score) * 100.0, 100.0)), 2)
    except (TypeError, ValueError):
        return None


def _format_score(score):
    """Redondea la puntuación nativa para estabilizar la respuesta JSON."""

    if score is None:
        return None
    try:
        return round(float(score), 6)
    except (TypeError, ValueError):
        return None


def _record_from_qdrant_point(point):
    """Transforma un punto completo en el contrato común de la API."""

    payload = _point_attr(point, 'payload', {}) or {}
    rid = _point_attr(point, 'id')
    raw_score = _point_attr(point, 'score')
    score = _format_score(raw_score)
    title = payload.get('title') or payload.get('name') or str(rid)
    return {
        'id': rid,
        'title': title,
        'type': payload.get('type') or payload.get('style_name'),
        'resolution': payload.get('resolution') or '',
        'base64': normalize_base64_payload(payload.get('base64')),
        'score': score,
        'score_percent': _score_percent(score)
    }


def _catalog_record_from_qdrant_point(point):
    """Transforma un punto del catálogo sin incluir datos de similaridad."""

    payload = _point_attr(point, 'payload', {}) or {}
    rid = _point_attr(point, 'id')
    return {
        'id': rid,
        'title': payload.get('title') or payload.get('name') or str(rid),
        'type': payload.get('type'),
        'resolution': payload.get('resolution') or '',
        'base64': normalize_base64_payload(payload.get('base64')),
    }


def _get_filtered_qdrant_records(q: QdrantClient, limit: int, type_filter=None, search=None):
    """Filtra metadatos primero y recupera payloads pesados solo para coincidencias."""

    if limit <= 0:
        return []

    selected_ids = []
    offset = None
    search_norm = _normalize_search_text(search)
    metadata_fields = ['title', 'name', 'type', 'resolution']

    while len(selected_ids) < limit:
        records, next_offset = q.scroll(
            collection_name=QDRANT_COLLECTION,
            with_payload=metadata_fields,
            with_vectors=False,
            limit=256,
            offset=offset,
        )
        for record in records:
            payload = _point_attr(record, 'payload', {}) or {}
            rid = _point_attr(record, 'id')
            title = payload.get('title') or payload.get('name') or str(rid)
            record_type = payload.get('type')
            if type_filter and record_type != type_filter:
                continue
            if search_norm and search_norm not in _normalize_search_text(title) and search_norm not in _normalize_search_text(record_type):
                continue
            if rid is not None:
                selected_ids.append(rid)
            if len(selected_ids) >= limit:
                break

        if len(selected_ids) >= limit or next_offset is None or next_offset == offset:
            break
        offset = next_offset

    if not selected_ids:
        return []

    full_records = q.retrieve(
        collection_name=QDRANT_COLLECTION,
        ids=selected_ids,
        with_payload=True,
        with_vectors=False,
    )
    records_by_id = {str(_point_attr(record, 'id')): record for record in full_records}
    ordered_records = [records_by_id[str(rid)] for rid in selected_ids if str(rid) in records_by_id]
    if len(ordered_records) != len(selected_ids):
        raise RuntimeError('Qdrant no devolvio todos los registros seleccionados.')
    return [_catalog_record_from_qdrant_point(record) for record in ordered_records]


def _normalize_vector(vector):
    """Obtiene una lista numérica incluso cuando Qdrant devuelve vectores nombrados."""

    if isinstance(vector, dict):
        for candidate in vector.values():
            normalized = _normalize_vector(candidate)
            if normalized is not None:
                return normalized
        return None
    if isinstance(vector, (list, tuple)):
        return list(vector)
    return None


def _retrieve_qdrant_point(
    q: QdrantClient,
    pid,
    with_vectors=False,
    with_payload=True,
    return_status=False,
    allow_scroll_fallback=True,
):
    """Busca un punto por ID y admite clientes con diferencias de serialización."""

    candidates = [pid]
    request_succeeded = False
    if not isinstance(pid, str):
        candidates.append(str(pid))
    for candidate in candidates:
        try:
            records = q.retrieve(
                collection_name=QDRANT_COLLECTION,
                ids=[candidate],
                with_payload=with_payload,
                with_vectors=with_vectors
            )
            request_succeeded = True
            points = _iter_qdrant_points(records)
            if points:
                return (points[0], True) if return_status else points[0]
        except Exception:
            continue
    if not allow_scroll_fallback:
        return (None, request_succeeded) if return_status else None
    try:
        records, _ = q.scroll(
            collection_name=QDRANT_COLLECTION,
            with_payload=with_payload,
            with_vectors=with_vectors,
            limit=10000
        )
        request_succeeded = True
        for point in records:
            if _ids_match(_point_attr(point, 'id'), pid):
                return (point, True) if return_status else point
    except Exception:
        pass
    return (None, request_succeeded) if return_status else None


def _qdrant_result_records(points, pid, limit, type_filter=None):
    """Excluye el origen, aplica el estilo y conserva el orden por puntuación."""

    out = []
    for point in points:
        rid = _point_attr(point, 'id')
        if _ids_match(rid, pid):
            continue
        payload = _point_attr(point, 'payload', {}) or {}
        if type_filter and payload.get('type') != type_filter:
            continue
        record = _record_from_qdrant_point(point)
        out.append(record)
        if len(out) >= limit:
            break
    return out


def _similarity_query_limit(q: QdrantClient, limit: int, type_filter=None):
    """Amplía la consulta cuando el filtro de estilo se aplica en la aplicación."""

    if not type_filter:
        return limit + 1
    try:
        count_result = q.count(collection_name=QDRANT_COLLECTION, exact=True)
        collection_count = int(getattr(count_result, 'count', count_result))
        return max(limit + 1, collection_count)
    except Exception:
        return 10000


def _hydrate_qdrant_records(q: QdrantClient, records):
    """Completa los payloads filtrados sin perder el orden ni la puntuación."""

    if not records:
        return []
    ids = [record['id'] for record in records]
    full_points = q.retrieve(
        collection_name=QDRANT_COLLECTION,
        ids=ids,
        with_payload=True,
        with_vectors=False,
    )
    points_by_id = {str(_point_attr(point, 'id')): point for point in full_points}
    hydrated = []
    for record in records:
        point = points_by_id.get(str(record['id']))
        if point is None:
            raise RuntimeError(f"Qdrant no devolvio el registro {record['id']}.")
        hydrated_record = _record_from_qdrant_point(point)
        hydrated_record['score'] = record.get('score')
        hydrated_record['score_percent'] = record.get('score_percent')
        hydrated.append(hydrated_record)
    return hydrated


def _query_qdrant_neighbors(q: QdrantClient, query, pid, limit: int, type_filter=None):
    """Ejecuta una consulta de vecinos por ID o por vector."""

    query_limit = _similarity_query_limit(q, limit, type_filter)
    with_payload = ['title', 'name', 'type', 'resolution'] if type_filter else True
    result = q.query_points(
        collection_name=QDRANT_COLLECTION,
        query=query,
        with_payload=with_payload,
        limit=query_limit,
    )
    records = _qdrant_result_records(
        _iter_qdrant_points(result),
        pid,
        limit,
        type_filter=type_filter,
    )
    return _hydrate_qdrant_records(q, records) if type_filter else records


def _get_qdrant_similar_records(
    id: str,
    limit: int = 12,
    strict: bool = False,
    required: bool = False,
    type_filter=None,
):
    """Resuelve similares con reintento vectorial y errores según el modo solicitado."""

    q = _get_qdrant_client()
    if q is None:
        if strict or required:
            raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para obtener la similaridad.')
        return None

    pid = _coerce_qdrant_id(id)
    errors = []
    similarity_request_succeeded = False
    primary_request_succeeded = False
    best_partial = []

    try:
        records = _query_qdrant_neighbors(q, pid, pid, limit, type_filter=type_filter)
        primary_request_succeeded = True
        similarity_request_succeeded = True
        best_partial = records
        if records and (not strict or (len(records) >= limit and all(item.get('score') is not None for item in records))):
            return records
    except Exception as exc:
        errors.append(str(exc))

    source = None
    source_lookup_succeeded = False
    try:
        source, source_lookup_succeeded = _retrieve_qdrant_point(
            q,
            pid,
            with_vectors=True,
            with_payload=False,
            return_status=True,
            allow_scroll_fallback=False,
        )
    except Exception as exc:
        errors.append(str(exc))

    if (strict or required) and source_lookup_succeeded and source is None and not best_partial:
        raise HTTPException(status_code=404, detail='La imagen consultada no existe en Qdrant.')

    vector = _normalize_vector(_point_attr(source, 'vector')) if source is not None else None
    should_query_by_vector = vector is not None and (not primary_request_succeeded or strict)
    if should_query_by_vector:
        try:
            records = _query_qdrant_neighbors(q, vector, pid, limit, type_filter=type_filter)
            similarity_request_succeeded = True
            if len(records) > len(best_partial):
                best_partial = records
            if records and (not strict or (len(records) >= limit and all(item.get('score') is not None for item in records))):
                return records
        except Exception as exc:
            errors.append(str(exc))

    if strict or required:
        if not similarity_request_succeeded:
            raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para obtener la similaridad.')
    if required:
        if not best_partial and not source_lookup_succeeded:
            raise HTTPException(status_code=503, detail='No se pudo verificar la imagen consultada en Qdrant.')
        return best_partial
    if strict:
        detail = f'Qdrant no devolvio {limit} resultados con score nativo para la imagen seleccionada.'
        if errors:
            detail += f' Ultimo error: {errors[-1]}'
        raise HTTPException(status_code=500, detail=detail)
    return best_partial or None


def _get_qdrant_record(id: str, strict: bool = False):
    """Recupera un único registro y distingue indisponibilidad de inexistencia."""

    q = _get_qdrant_client()
    if q is None:
        if strict:
            raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para obtener la imagen consultada.')
        return None
    pid = _coerce_qdrant_id(id)
    point, request_succeeded = _retrieve_qdrant_point(q, pid, with_vectors=False, return_status=True)
    if point is None:
        if strict:
            if not request_succeeded:
                raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para obtener la imagen consultada.')
            raise HTTPException(status_code=404, detail='La imagen consultada no existe en Qdrant.')
        return None
    record = _record_from_qdrant_point(point)
    record['score'] = None
    record['score_percent'] = None
    return record


# Metadatos relacionales y persistencia del historial de exportaciones.
def _open_mysql_connection_or_error():
    """Abre MySQL o convierte el fallo de configuración/conexión en un error HTTP."""

    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        raise HTTPException(status_code=500, detail='Configuracion de MySQL incompleta.')
    try:
        return mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
    except Error as exc:
        raise HTTPException(status_code=500, detail=f'No se pudo conectar con MySQL: {exc}')


def _fetch_artworks_by_qdrant_ids(ids, strict=True):
    """Obtiene de MySQL los metadatos asociados a una lista de puntos de Qdrant."""

    unique_ids = []
    for value in ids:
        text = str(value)
        if text not in unique_ids:
            unique_ids.append(text)
    if not unique_ids:
        return {}

    try:
        conn = _open_mysql_connection_or_error()
    except HTTPException as exc:
        if strict:
            raise
        return {}, [f"No se pudo consultar MySQL: {exc.detail}"]
    cursor = conn.cursor(dictionary=True)
    try:
        placeholders = ','.join(['%s'] * len(unique_ids))
        cursor.execute(
            "SELECT title, author_name, style_name, source_name, source_url, file_path, id_qdrant_point "
            f"FROM artworks WHERE id_qdrant_point IN ({placeholders})",
            tuple(unique_ids)
        )
        rows = cursor.fetchall() or []
    except Error as exc:
        if strict:
            raise HTTPException(status_code=500, detail=f'No se pudieron consultar los metadatos de artworks: {exc}')
        return {}, [f"No se pudieron consultar los metadatos de artworks: {exc}"]
    finally:
        cursor.close()
        conn.close()

    found = {str(row.get('id_qdrant_point')): row for row in rows}
    missing = [value for value in unique_ids if value not in found]
    if missing and strict:
        raise HTTPException(
            status_code=500,
            detail='Faltan metadatos en artworks para id_qdrant_point: ' + ', '.join(missing)
        )
    return found, []


def _fallback_artwork_from_record(record):
    """Crea metadatos mínimos cuando la tabla artworks no contiene un punto."""

    return {
        'title': record.get('title') or str(record.get('id')),
        'author_name': None,
        'style_name': record.get('type'),
        'source_name': None,
        'source_url': None,
        'file_path': '',
        'id_qdrant_point': str(record.get('id')),
    }


def _artworks_with_warnings(records, fetched_artworks):
    """Combina metadatos disponibles y enumera los campos faltantes del reporte."""

    required_fields = ('title', 'author_name', 'style_name', 'source_name', 'source_url')
    artworks = {}
    warnings = []
    for index, record in enumerate(records):
        rid = str(record.get('id'))
        artwork = dict(fetched_artworks.get(rid) or _fallback_artwork_from_record(record))
        if not artwork.get('title'):
            artwork['title'] = record.get('title') or rid
        if not artwork.get('style_name'):
            artwork['style_name'] = record.get('type')
        missing_fields = [field for field in required_fields if not artwork.get(field)]
        if missing_fields:
            location = 'la Imagen consultada' if index == 0 else f'el Resultado #{index}'
            warnings.append(
                f"Metadatos incompletos en {location}: faltan {', '.join(missing_fields)}."
            )
        artworks[rid] = artwork
    return artworks, warnings


def _insert_exported_file(report_name, report_path):
    """Guarda el DOCX y conserva únicamente las diez exportaciones más recientes."""

    try:
        file_data = report_path.read_bytes()
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f'No se pudo leer el reporte generado: {exc}')

    conn = _open_mysql_connection_or_error()
    cursor = conn.cursor(dictionary=True)
    expired_paths = []
    try:
        conn.start_transaction()
        cursor.execute(
            "INSERT INTO exported_files (report_name, file_format, file_path, created_at, `file`) "
            "VALUES (%s, %s, %s, CURDATE(), %s)",
            (report_name, 'docx', str(report_path), file_data)
        )
        cursor.execute(
            "SELECT id_export, file_path FROM exported_files "
            "ORDER BY id_export DESC LIMIT 10, 18446744073709551615 FOR UPDATE"
        )
        expired = cursor.fetchall() or []
        if expired:
            expired_ids = [int(row['id_export']) for row in expired]
            placeholders = ','.join(['%s'] * len(expired_ids))
            cursor.execute(
                f"DELETE FROM exported_files WHERE id_export IN ({placeholders})",
                tuple(expired_ids),
            )
            expired_paths = [row.get('file_path') for row in expired]
        conn.commit()
    except Error as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f'No se pudo registrar el reporte exportado: {exc}')
    finally:
        cursor.close()
        conn.close()

    reports_root = REPORTS_DIR.resolve()
    for expired_path in expired_paths:
        if not expired_path:
            continue
        candidate = Path(expired_path)
        try:
            resolved = candidate.resolve()
            resolved.relative_to(reports_root)
            if resolved.is_file():
                resolved.unlink()
        except (OSError, ValueError):
            continue


# Composición del documento DOCX y de sus recursos gráficos.
def _resolve_artwork_image_path(artwork):
    """Resuelve las variantes históricas de rutas almacenadas en artworks."""

    file_path = str(artwork.get('file_path') or '').replace('\\', '/')
    candidates = []
    if file_path:
        raw = Path(file_path)
        if raw.is_absolute():
            candidates.append(raw)
        candidates.append(ROOT / file_path)
        if file_path.startswith('ArtStylesDataset/'):
            candidates.append(ROOT / 'ipynb' / file_path)
        candidates.append(DATASET_DIR / file_path)
    for candidate in candidates:
        try:
            if candidate.exists() and candidate.is_file():
                return candidate
        except Exception:
            continue
    return None


def _image_stream_for_record(record, artwork):
    """Prioriza la imagen de Qdrant y usa el dataset local como respaldo del reporte."""

    stream = base64_to_image_stream(record.get('base64'))
    if stream is not None:
        return stream
    path = _resolve_artwork_image_path(artwork)
    if path is not None:
        return image_to_bytes(path)
    raise HTTPException(status_code=500, detail=f"No se pudo cargar la imagen para Qdrant id {record.get('id')}.")


def _clear_docx_body(document):
    """Vacía la plantilla conservando sus propiedades de sección."""

    body = document._element.body
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)


def _meta_value(artwork, key):
    """Presenta un valor legible cuando faltan metadatos."""

    value = artwork.get(key)
    return str(value) if value not in (None, '') else 'No disponible'


def _add_artwork_table(document, artwork, score_percent=None):
    """Agrega una tabla de metadatos a dos columnas."""

    rows = [
        ('Titulo', _meta_value(artwork, 'title')),
        ('Autor', _meta_value(artwork, 'author_name')),
        ('Estilo', _meta_value(artwork, 'style_name')),
        ('Fuente', _meta_value(artwork, 'source_name')),
        ('URL', _meta_value(artwork, 'source_url')),
    ]
    if score_percent is not None:
        rows.append(('Score', f'{score_percent:.2f}%'))
    table = document.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
    return table


def _set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    """Configura márgenes internos en unidades DXA de Word."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def _keep_table_row_together(row):
    """Evita que Word divida una fila entre dos páginas."""

    from docx.oxml import OxmlElement

    row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))


def _add_fitted_picture(run, image_stream, max_width, max_height):
    """Inserta una imagen ajustada al marco sin deformar su proporción."""

    from docx.shared import Inches

    image_stream.seek(0)
    with Image.open(image_stream) as image:
        pixel_width, pixel_height = image.size
    image_stream.seek(0)
    if pixel_width <= 0 or pixel_height <= 0:
        raise HTTPException(status_code=500, detail='No se pudo determinar el tamano de una imagen del reporte.')

    aspect_ratio = pixel_width / pixel_height
    if aspect_ratio >= max_width / max_height:
        width = max_width
        height = max_width / aspect_ratio
    else:
        height = max_height
        width = max_height * aspect_ratio
    return run.add_picture(image_stream, width=Inches(width), height=Inches(height))


def _add_compact_result(document, index, record, artwork):
    """Compone una fila compacta con miniatura, metadatos y puntuación."""

    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    heading = document.add_heading(f'Resultado #{index}', level=2)
    heading.paragraph_format.keep_with_next = True

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.25)
    image_cell, metadata_cell = table.rows[0].cells
    image_cell.width = Inches(1.75)
    metadata_cell.width = Inches(4.25)
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    metadata_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(image_cell)
    _set_cell_margins(metadata_cell)
    _keep_table_row_together(table.rows[0])

    image_paragraph = image_cell.paragraphs[0]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_after = Pt(0)
    _add_fitted_picture(
        image_paragraph.add_run(),
        _image_stream_for_record(record, artwork),
        max_width=1.5,
        max_height=1.65,
    )

    rows = [
        ('Titulo', _meta_value(artwork, 'title')),
        ('Autor', _meta_value(artwork, 'author_name')),
        ('Estilo', _meta_value(artwork, 'style_name')),
        ('Fuente', _meta_value(artwork, 'source_name')),
        ('URL', _meta_value(artwork, 'source_url')),
        ('Score', f"{float(record['score_percent']):.2f}%"),
    ]
    metadata_cell.text = ''
    for row_index, (label, value) in enumerate(rows):
        paragraph = metadata_cell.paragraphs[0] if row_index == 0 else metadata_cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        label_run = paragraph.add_run(f'{label}: ')
        label_run.bold = True
        label_run.font.size = Pt(8.5)
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(8.5)


def _add_style_count_table(document, style_counts):
    """Distribuye el resumen de estilos en dos pares de columnas."""

    from docx.shared import Pt

    entries = list(style_counts.items())
    row_count = (len(entries) + 1) // 2
    table = document.add_table(rows=row_count, cols=4)
    table.style = 'Table Grid'
    for row_index, row in enumerate(table.rows):
        _keep_table_row_together(row)
        for cell in row.cells:
            _set_cell_margins(cell, top=35, start=60, bottom=35, end=60)
        for pair_index in range(2):
            entry_index = row_index + pair_index * row_count
            if entry_index >= len(entries):
                continue
            style, count = entries[entry_index]
            style_cell = row.cells[pair_index * 2]
            count_cell = row.cells[pair_index * 2 + 1]
            style_cell.text = str(style)
            count_cell.text = str(count)
            for cell in (style_cell, count_cell):
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.runs[0].font.size = Pt(8.5)
    return table


def _make_score_chart(results):
    """Genera en memoria el gráfico horizontal de similaridad."""

    import matplotlib
    matplotlib.use('Agg', force=True)
    import matplotlib.pyplot as plt

    labels = []
    values = []
    for idx, item in enumerate(results, start=1):
        labels.append(f'Resultado #{idx}')
        values.append(float(item['score_percent']))

    fig_height = max(4.5, len(results) * 0.36)
    fig, ax = plt.subplots(figsize=(8.6, fig_height))
    y_pos = list(range(len(values)))
    ax.barh(y_pos, values, color='#B39DDB')
    ax.set_yticks(y_pos, labels)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel('Score de similaridad (%)')
    ax.grid(axis='x', alpha=0.25)
    for y, value in zip(y_pos, values):
        ax.text(min(value + 1, 98), y, f'{value:.2f}%', va='center', fontsize=8)
    fig.tight_layout()
    stream = io.BytesIO()
    fig.savefig(stream, format='png', dpi=160)
    plt.close(fig)
    stream.seek(0)
    return stream


def _make_style_chart(style_counts):
    """Genera en memoria el gráfico circular de estilos recuperados."""

    import matplotlib
    matplotlib.use('Agg', force=True)
    import matplotlib.pyplot as plt

    labels = list(style_counts.keys())
    values = list(style_counts.values())
    colors = [_style_chart_color(label) for label in labels]
    fig, ax = plt.subplots(figsize=(6.2, 4.8))
    ax.pie(
        values,
        labels=labels,
        colors=colors,
        autopct='%1.0f%%',
        startangle=90,
        wedgeprops={'edgecolor': 'white', 'linewidth': 1},
    )
    ax.axis('equal')
    fig.tight_layout()
    stream = io.BytesIO()
    fig.savefig(stream, format='png', dpi=160)
    plt.close(fig)
    stream.seek(0)
    return stream


def _style_chart_color(style_name):
    """Asigna colores estables a estilos conocidos y desconocidos."""

    normalized = unicodedata.normalize('NFKD', str(style_name or ''))
    normalized = normalized.encode('ascii', 'ignore').decode('ascii').lower()
    normalized = re.sub(r'[^a-z0-9]+', '_', normalized).strip('_')
    if normalized in STYLE_CHART_COLORS:
        return STYLE_CHART_COLORS[normalized]
    digest = hashlib.sha256(normalized.encode('utf-8')).digest()
    return STYLE_CHART_FALLBACK_COLORS[digest[0] % len(STYLE_CHART_FALLBACK_COLORS)]


def _build_similarity_docx(
    query_record,
    similar_records,
    artworks,
    output_path,
    metadata_warnings=None,
    report_warnings=None,
):
    """Construye el informe completo a partir de la plantilla institucional."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    if not REPORT_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail='No se encontro la plantilla DOCX del reporte.')

    document = Document(str(REPORT_TEMPLATE_PATH))
    _clear_docx_body(document)
    pastel_purple = RGBColor(0xB3, 0x9D, 0xDB)
    for style_name in ('Heading 1', 'Heading 2'):
        document.styles[style_name].font.color.rgb = pastel_purple
    document.styles['Heading 1'].paragraph_format.space_before = Pt(10)
    document.styles['Heading 1'].paragraph_format.space_after = Pt(4)
    document.styles['Heading 2'].paragraph_format.space_before = Pt(5)
    document.styles['Heading 2'].paragraph_format.space_after = Pt(2)

    query_artwork = artworks[str(query_record['id'])]
    style_counts = Counter(_meta_value(artworks[str(item['id'])], 'style_name') for item in similar_records)

    title = document.add_heading('CONSULTA DE SIMILARIDAD', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        'La evidencia de comparacion se fundamenta en representaciones vectoriales generadas mediante ResNet '
        'y busqueda de vecinos cercanos utilizando Qdrant.'
    )
    if metadata_warnings:
        document.add_heading('ADVERTENCIAS DE METADATOS', level=1)
        document.add_paragraph(
            'El reporte se genero de forma parcial. Algunos registros no tienen todos los metadatos esperados en artworks.'
        )
        for warning in metadata_warnings:
            document.add_paragraph(f'- {warning}')
    if report_warnings:
        document.add_heading('ADVERTENCIAS DEL REPORTE', level=1)
        for warning in report_warnings:
            document.add_paragraph(f'- {warning}')

    document.add_heading('IMAGEN CONSULTADA', level=1)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(_image_stream_for_record(query_record, query_artwork), width=Inches(4.2))
    _add_artwork_table(document, query_artwork)

    document.add_heading('RESUMEN DEL ANALISIS', level=1)
    document.add_paragraph(f'Total de imagenes comparadas: {len(similar_records)}')
    document.add_paragraph('Metodo de comparacion: Qdrant sobre embeddings de imagen.')

    document.add_heading('ANALISIS DE SIMILARIDAD', level=1)
    document.add_paragraph(
        'Los valores mostrados representan el grado de proximidad vectorial respecto a la imagen consultada.'
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(_make_score_chart(similar_records), width=Inches(6.6))

    document.add_page_break()
    results_heading = document.add_heading('IMAGENES SIMILARES RECUPERADAS', level=1)
    results_heading.paragraph_format.keep_with_next = True
    for index, item in enumerate(similar_records, start=1):
        artwork = artworks[str(item['id'])]
        _add_compact_result(document, index, item, artwork)

    document.add_page_break()
    styles_heading = document.add_heading('DISTRIBUCION DE ESTILOS RECUPERADOS', level=1)
    styles_heading.paragraph_format.keep_with_next = True
    _add_style_count_table(document, style_counts)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(_make_style_chart(style_counts), width=Inches(5.0))

    document.save(str(output_path))


def _next_similarity_report_path(id, report_date=None):
    """Encuentra un nombre disponible sin sobrescribir reportes existentes."""

    report_date = report_date or datetime.now()
    date_stamp = report_date.strftime('%d-%m-%Y')
    safe_id = _safe_filename_part(id)
    report_stem = f'consulta_similaridad_{safe_id}_{date_stamp}'
    report_name = f'{report_stem}.docx'
    report_path = REPORTS_DIR / report_name
    suffix = 2
    while report_path.exists():
        report_name = f'{report_stem}_{suffix}.docx'
        report_path = REPORTS_DIR / report_name
        suffix += 1
    return report_path, report_name


def _generate_similarity_report(id: str, limit: int = 12):
    """Coordina datos, documento, persistencia y limpieza ante errores."""

    if limit != 12:
        limit = 12
    query_record = _get_qdrant_record(id, strict=True)
    similar_records = _get_qdrant_similar_records(id, limit=limit, strict=True)
    if len(similar_records) != limit:
        raise HTTPException(status_code=500, detail=f'Qdrant devolvio {len(similar_records)} resultados; se requieren {limit}.')
    if any(item.get('score') is None or item.get('score_percent') is None for item in similar_records):
        raise HTTPException(status_code=500, detail='Qdrant no devolvio scores nativos para todos los resultados.')

    all_ids = [query_record['id']] + [item['id'] for item in similar_records]
    fetched_artworks, fetch_warnings = _fetch_artworks_by_qdrant_ids(all_ids, strict=False)
    artworks, metadata_warnings = _artworks_with_warnings([query_record] + similar_records, fetched_artworks)

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    report_path, report_name = _next_similarity_report_path(id)
    _build_similarity_docx(
        query_record,
        similar_records,
        artworks,
        report_path,
        metadata_warnings=metadata_warnings,
        report_warnings=fetch_warnings,
    )
    try:
        _insert_exported_file(report_name, report_path)
    except Exception:
        try:
            report_path.unlink(missing_ok=True)
        except OSError:
            pass
        raise
    return report_path, report_name


# API pública del catálogo y los reportes de similaridad.
@app.get("/api/types")
def get_types():
    """Lista los estilos presentes en la colección de Qdrant."""

    q = _get_qdrant_client()
    if q is None:
        raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para cargar los tipos.')
    try:
        records, _ = q.scroll(
            collection_name=QDRANT_COLLECTION,
            with_payload=['type'],
            with_vectors=False,
            limit=1000
        )
        types = set()
        for r in records:
            payload = getattr(r, 'payload', None)
            if isinstance(payload, dict):
                t = payload.get('type')
                if t:
                    types.add(t)
        return sorted(types)
    except Exception as exc:
        raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para cargar los tipos.') from exc


@app.get("/api/records")
def get_records(limit: int=12, type: str=None, search: str=None):
    """Devuelve registros del catálogo, opcionalmente filtrados por estilo o texto."""

    q = _get_qdrant_client()
    if q is None:
        raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para cargar las imagenes.')

    try:
        if type or search:
            return _get_filtered_qdrant_records(q, limit=limit, type_filter=type, search=search)
        recs, _ = q.scroll(
            collection_name=QDRANT_COLLECTION,
            with_payload=True,
            with_vectors=False,
            limit=limit,
        )
        return [_catalog_record_from_qdrant_point(record) for record in recs]
    except Exception as exc:
        raise HTTPException(status_code=503, detail='No se pudo conectar con Qdrant para cargar las imagenes.') from exc


@app.get("/api/records/{id:path}/similar")
def get_similar(id: str, limit: int=12, type: str=None):
    """Devuelve los vecinos visuales de un registro de Qdrant."""

    return _get_qdrant_similar_records(id, limit=limit, required=True, type_filter=type)


@app.get("/api/records/{id:path}/similar-report")
def get_similarity_report(id: str, limit: int=12):
    """Genera y descarga el informe DOCX de una consulta."""

    report_path, report_name = _generate_similarity_report(id, limit=limit)
    return FileResponse(
        report_path,
        media_type=DOCX_MEDIA_TYPE,
        filename=report_name
    )


@app.get("/api/records/{id:path}")
def get_record(id: str):
    """Devuelve una imagen local cuando su ruta relativa se solicita de forma explícita."""

    p = DATASET_DIR / id
    if not p.exists():
        raise HTTPException(404,"Record not found")
    b64,size = image_to_base64(p, max_size=(1600,1200))
    return {"id": id, "title": p.stem, "type": p.parent.name, "resolution": f"{size[0]} × {size[1]}", "base64": b64}


# Configuración, autenticación y operaciones administrativas.
def _load_mysql_config():
    """Combina la configuración de entorno y el archivo local de secretos."""

    # Las variables de entorno tienen prioridad sobre el archivo local.
    cfg = {
        'host': os.environ.get('MYSQL_HOST'),
        'port': os.environ.get('MYSQL_PORT'),
        'user': os.environ.get('MYSQL_USER'),
        'password': os.environ.get('MYSQL_PASSWORD'),
        'database': os.environ.get('MYSQL_DATABASE')
    }
    # Completa los valores ausentes desde frontend/.config/secrets.toml.
    if not cfg['user'] or cfg['database'] is None or cfg['password'] is None:
        secrets_path = ROOT / 'frontend' / '.config' / 'secrets.toml'
        if secrets_path.exists():
            try:
                data = toml.load(secrets_path)
                cfg['host'] = cfg['host'] or data.get('mysql_host')
                cfg['port'] = cfg['port'] or data.get('mysql_port')
                cfg['user'] = cfg['user'] or data.get('mysql_user')
                if cfg['password'] is None:
                    cfg['password'] = data.get('mysql_password')
                cfg['database'] = cfg['database'] or data.get('mysql_database')
            except Exception:
                pass
    # Normaliza el host y el puerto antes de abrir cualquier conexión.
    if cfg['host'] is None:
        cfg['host'] = '127.0.0.1'
    try:
        cfg['port'] = int(cfg['port']) if cfg['port'] else 3306
    except Exception:
        cfg['port'] = 3306
    return cfg


def _get_qdrant_client() -> typing.Optional[QdrantClient]:
    """Crea un cliente de Qdrant solo cuando existen URL y clave válidas."""

    # Las variables de entorno tienen prioridad sobre el archivo local.
    url = os.environ.get('QDRANT_URL') or os.environ.get('QDRANT_DB_URL')
    key = os.environ.get('QDRANT_API_KEY') or os.environ.get('QDRANT_APIKEY')
    # Completa las credenciales desde el archivo local de secretos.
    if not url or not key:
        secrets_path = ROOT / 'frontend' / '.config' / 'secrets.toml'
        try:
            if secrets_path.exists():
                data = toml.load(secrets_path)
                url = url or data.get('qdrant_db_url')
                key = key or data.get('qdrant_api_key')
        except Exception:
            pass
    if not url or not key:
        return None
    try:
        return QdrantClient(url=url, api_key=key)
    except Exception:
        return None


def _normalize_email(email: str) -> str:
    """Normaliza el correo antes de validarlo o consultarlo."""

    if not isinstance(email, str):
        return ''
    return email.strip().lower()


def _is_valid_email(email: str) -> bool:
    """Aplica la validación estructural y el límite aceptado por la interfaz."""

    if not email or len(email) > 128:
        return False
    return bool(re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email))


def _get_admin_from_request(request: Request):
    """Exige una cookie vinculada a un administrador activo."""

    email = request.cookies.get('admin_email')
    if not email:
        raise HTTPException(status_code=401, detail='Not authenticated')
    admin = fetch_admin_by_email(email)
    if not admin or not bool(admin.get('is_active')):
        raise HTTPException(status_code=401, detail='Not authenticated')
    return admin


def _fetch_exported_files_history():
    """Obtiene las diez exportaciones más recientes sin cargar sus BLOB."""

    conn = _open_mysql_connection_or_error()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT id_export, report_name, file_format, created_at, OCTET_LENGTH(`file`) AS size_bytes "
            "FROM exported_files ORDER BY id_export DESC LIMIT 10"
        )
        rows = cursor.fetchall() or []
        for row in rows:
            created_at = row.get('created_at')
            row['created_at'] = created_at.isoformat() if hasattr(created_at, 'isoformat') else str(created_at or '')
            row['size_bytes'] = int(row.get('size_bytes') or 0)
        return rows
    except Error as exc:
        raise HTTPException(status_code=500, detail=f'No se pudo cargar el historial de reportes: {exc}')
    finally:
        cursor.close()
        conn.close()


def _fetch_exported_file_content(export_id):
    """Recupera el BLOB de una exportación para su descarga."""

    conn = _open_mysql_connection_or_error()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT report_name, file_format, `file` FROM exported_files WHERE id_export = %s LIMIT 1",
            (export_id,),
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail='Reporte no encontrado')
        row['file'] = bytes(row.get('file') or b'')
        if not row['file']:
            raise HTTPException(status_code=404, detail='El reporte no tiene un archivo almacenado')
        return row
    except Error as exc:
        raise HTTPException(status_code=500, detail=f'No se pudo descargar el reporte: {exc}')
    finally:
        cursor.close()
        conn.close()


@app.get('/api/exported-files')
def api_get_exported_files(request: Request):
    """Expone el historial de exportaciones a administradores autenticados."""

    _get_admin_from_request(request)
    return JSONResponse(
        content=_fetch_exported_files_history(),
        headers={'Cache-Control': 'no-store'},
    )


@app.get('/api/exported-files/{export_id}/download')
def api_download_exported_file(export_id: int, request: Request):
    """Descarga una exportación almacenada en MySQL."""

    _get_admin_from_request(request)
    exported_file = _fetch_exported_file_content(export_id)
    report_name = Path(str(exported_file['report_name'])).name
    safe_name = re.sub(r'[^a-zA-Z0-9_.-]+', '_', report_name) or f'reporte_{export_id}.docx'
    return Response(
        content=exported_file['file'],
        media_type=DOCX_MEDIA_TYPE,
        headers={
            'Content-Disposition': f'attachment; filename="{safe_name}"',
            'Cache-Control': 'no-store',
        },
    )


# Mensajes de contacto y bandeja administrativa.
def _save_contact_message(sender_name: str, sender_email: str, message_body: str):
    """Persiste un mensaje validado desde la página informativa."""

    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        raise HTTPException(status_code=500, detail='Database configuration missing')
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor()
        try:
            cursor.execute(
                "INSERT INTO contact_messages (sender_name, sender_email, message_body, created_at, is_reviwed) "
                "VALUES (%s, %s, %s, NOW(), FALSE)",
                (sender_name, sender_email, message_body)
            )
            conn.commit()
            return cursor.lastrowid
        finally:
            cursor.close()
    except Error:
        raise HTTPException(status_code=500, detail='Failed to save contact message')


def _fetch_contact_messages():
    """Lista los mensajes de contacto del más reciente al más antiguo."""

    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        raise HTTPException(status_code=500, detail='Database configuration missing')
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor(dictionary=True)
        try:
            cursor.execute(
                "SELECT id_message, sender_name, sender_email, message_body, created_at, is_reviwed "
                "FROM contact_messages ORDER BY created_at DESC"
            )
            rows = cursor.fetchall() or []
            for row in rows:
                created_at = row.get('created_at')
                if hasattr(created_at, 'isoformat'):
                    try:
                        row['created_at'] = created_at.isoformat(' ')
                    except TypeError:
                        row['created_at'] = created_at.isoformat()
                else:
                    row['created_at'] = str(created_at)
                row['is_reviwed'] = bool(row.get('is_reviwed'))
            return rows
        finally:
            cursor.close()
    except Error:
        raise HTTPException(status_code=500, detail='Failed to load contact messages')


def _update_contact_message_reviewed(message_id: int, reviewed: bool):
    """Cambia el estado de revisión de un mensaje."""

    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        raise HTTPException(status_code=500, detail='Database configuration missing')
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE contact_messages SET is_reviwed = %s WHERE id_message = %s",
                (1 if reviewed else 0, message_id)
            )
            conn.commit()
            return cursor.rowcount > 0
        finally:
            cursor.close()
    except Error:
        raise HTTPException(status_code=500, detail='Failed to update message')


def _delete_contact_message(message_id: int):
    """Elimina un mensaje de la bandeja administrativa."""

    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        raise HTTPException(status_code=500, detail='Database configuration missing')
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor()
        try:
            cursor.execute(
                "DELETE FROM contact_messages WHERE id_message = %s",
                (message_id,)
            )
            conn.commit()
            return cursor.rowcount > 0
        finally:
            cursor.close()
    except Error:
        raise HTTPException(status_code=500, detail='Failed to delete message')


@app.get('/api/messages')
def api_get_messages(request: Request):
    """Entrega la bandeja solo a administradores autenticados."""

    _get_admin_from_request(request)
    return _fetch_contact_messages()


@app.patch('/api/messages/{message_id}/reviewed')
async def api_review_message(message_id: int, request: Request):
    """Valida y actualiza el indicador de revisión."""

    _get_admin_from_request(request)
    try:
        payload = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail='Invalid JSON payload')
    reviewed = payload.get('reviewed')
    if reviewed is None:
        raise HTTPException(status_code=400, detail='Missing reviewed state')
    if not isinstance(reviewed, bool):
        raise HTTPException(status_code=400, detail='Reviewed must be a boolean')
    if not _update_contact_message_reviewed(message_id, reviewed):
        raise HTTPException(status_code=404, detail='Mensaje no encontrado')
    return {'ok': True, 'reviewed': reviewed}


@app.delete('/api/messages/{message_id}')
def api_delete_message(message_id: int, request: Request):
    """Elimina un mensaje después de comprobar la sesión."""

    _get_admin_from_request(request)
    if not _delete_contact_message(message_id):
        raise HTTPException(status_code=404, detail='Mensaje no encontrado')
    return {'ok': True}


@app.post('/api/contact')
async def api_contact(request: Request):
    """Valida los límites del formulario y guarda el mensaje."""

    try:
        payload = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail='Invalid JSON payload')

    sender_name = payload.get('sender_name', '')
    sender_email = _normalize_email(payload.get('sender_email', ''))
    message_body = payload.get('message_body', '')

    if not isinstance(sender_name, str) or not sender_name.strip():
        raise HTTPException(status_code=400, detail='El nombre es obligatorio')
    sender_name = sender_name.strip()
    if len(sender_name) < 2 or len(sender_name) > 100:
        raise HTTPException(status_code=400, detail='El nombre debe tener entre 2 y 100 caracteres')

    if not _is_valid_email(sender_email):
        raise HTTPException(status_code=400, detail='Correo electrónico inválido')

    if not isinstance(message_body, str) or not message_body.strip():
        raise HTTPException(status_code=400, detail='El mensaje es obligatorio')
    message_body = message_body.strip()
    if len(message_body) < 10 or len(message_body) > 800:
        raise HTTPException(status_code=400, detail='El mensaje debe tener entre 10 y 800 caracteres')

    message_id = _save_contact_message(sender_name, sender_email, message_body)
    return {'ok': True, 'message_id': message_id}


def _verify_password(password: str, stored_password: typing.Optional[str]) -> bool:
    """Acepta el hash SHA-256 existente y mantiene compatibilidad con datos heredados."""

    if not isinstance(password, str) or not stored_password:
        return False
    if isinstance(stored_password, str) and len(stored_password) == 64 and re.fullmatch(r"[0-9a-fA-F]{64}", stored_password):
        return hashlib.sha256(password.encode("utf-8")).hexdigest().lower() == stored_password.lower()
    return password == str(stored_password)


def _update_admin_active_status(email: str, active: bool):
    """Sincroniza el indicador de sesión usado por las vistas administrativas."""

    if not email:
        return False
    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        return False
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE admins SET is_active = %s WHERE email_admin = %s",
                (1 if active else 0, email)
            )
            conn.commit()
            return cursor.rowcount > 0
        except Error:
            return False
        finally:
            cursor.close()
    except Error:
        return False
    finally:
        try:
            conn.close()
        except Exception:
            pass


def fetch_admin_by_email(email: str):
    """Carga un administrador y admite ambos nombres históricos de contraseña."""

    if not email:
        return None
    cfg = _load_mysql_config()
    if not cfg['user'] or cfg['database'] is None:
        return None
    try:
        conn = mysql.connector.connect(
            host=cfg['host'] or '127.0.0.1',
            port=cfg['port'],
            user=cfg['user'],
            password=cfg['password'],
            database=cfg['database']
        )
        cursor = conn.cursor(dictionary=True)
        query_base = "SELECT id_admin, email_admin, name_admin, is_active, created_at"
        try:
            cursor.execute(
                query_base + ", password_hash FROM admins WHERE email_admin = %s LIMIT 1",
                (email,)
            )
        except Error as error:
            if 'Unknown column' in str(error) or 'password_hash' in str(error):
                cursor.execute(
                    query_base + ", password_admin FROM admins WHERE email_admin = %s LIMIT 1",
                    (email,)
                )
            else:
                raise
        rows = cursor.fetchall()
        admin = rows[0] if rows else None
        if admin is not None:
            if 'password_hash' not in admin:
                admin['password_hash'] = None
            if 'password_admin' not in admin:
                admin['password_admin'] = None
        return admin
    except Error:
        return None
    finally:
        try:
            cursor.close()
        except Exception:
            pass
        try:
            conn.close()
        except Exception:
            pass


@app.post('/api/login')
async def api_login(request: Request, response: Response):
    """Valida credenciales, activa la sesión y establece la cookie HTTP-only."""

    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail='Invalid JSON')
    email = _normalize_email(data.get('email'))
    password = data.get('password') or ''
    if not _is_valid_email(email):
        raise HTTPException(status_code=400, detail='Invalid email format')
    if not isinstance(password, str) or len(password) < 8 or len(password) > 64:
        raise HTTPException(status_code=400, detail='Invalid password length')
    admin = fetch_admin_by_email(email)
    if not admin:
        raise HTTPException(status_code=401, detail='Credenciales inválidas')
    stored_password = admin.get('password_hash') or admin.get('password_admin')
    if stored_password is None:
        raise HTTPException(status_code=500, detail='Admin password is not configured')
    if not _verify_password(password, stored_password):
        raise HTTPException(status_code=401, detail='Credenciales inválidas')
    _update_admin_active_status(email, True)
    response.set_cookie(key='admin_email', value=admin.get('email_admin'), httponly=True)
    return {'ok': True, 'is_admin': True, 'email': admin.get('email_admin'), 'name': admin.get('name_admin')}


@app.post('/api/logout')
def api_logout(request: Request, response: Response):
    """Desactiva la sesión persistida y elimina la cookie."""

    email = request.cookies.get('admin_email')
    if email:
        _update_admin_active_status(email, False)
    response.delete_cookie('admin_email')
    return {'ok': True}


@app.get('/api/me')
def api_me(request: Request):
    """Informa a la interfaz si la cookie representa una sesión activa."""

    email = request.cookies.get('admin_email')
    if not email:
        return {'is_admin': False}
    admin = fetch_admin_by_email(email)
    if not admin:
        return {'is_admin': False}
    return {'is_admin': bool(admin.get('is_active')), 'email': admin.get('email_admin'), 'name': admin.get('name_admin')}


@app.get('/api/qdrant_status')
def api_qdrant_status():
    """Devuelve el estado de la conexión a Qdrant y alguna info básica de la colección."""
    q = _get_qdrant_client()
    if not q:
        return {'connected': False, 'reason': 'credentials_missing'}
    try:
        # Obtiene metadatos cuando la versión instalada del cliente lo permite.
        try:
            info = q.get_collection(collection_name=QDRANT_COLLECTION)
        except Exception:
            # Algunas versiones de qdrant-client no exponen get_collection.
            info = None
        # El conteo se informa por separado porque también puede no estar disponible.
        points = None
        try:
            stats = q.count(collection_name=QDRANT_COLLECTION)
            points = getattr(stats, 'count', stats)
        except Exception:
            points = None
        return {'connected': True, 'collection': QDRANT_COLLECTION, 'collection_info': str(info), 'points_count': points}
    except Exception as e:
        return {'connected': False, 'error': str(e)}


@app.get('/api/qdrant_debug')
def api_qdrant_debug(id: str = "0"):
    """Diagnostica compatibilidad del cliente Qdrant sin alterar la colección."""

    q = _get_qdrant_client()
    if not q:
        return {'ok': False, 'reason': 'no_client'}
    out = {'id': id}
    pid = None
    try:
        pid = int(id) if id.isdigit() else id
    except Exception:
        pid = id
    # Prueba primero la operación histórica de recomendación por identificador.
    try:
        recs = q.recommend_points(collection_name=QDRANT_COLLECTION, positive=[pid], limit=10)
        out['recommend_points_type'] = type(recs).__name__
        # Extrae identificadores de las distintas formas de respuesta conocidas.
        ids = []
        try:
            for r in recs:
                rid = getattr(r, 'id', None)
                if rid is None and isinstance(r, dict):
                    rid = r.get('id')
                ids.append(rid)
        except Exception:
            # Algunas versiones envuelven la lista en los atributos points o result.
            try:
                for r in getattr(recs, 'points', []) or getattr(recs, 'result', []):
                    ids.append(getattr(r, 'id', None) or (r.get('id') if isinstance(r, dict) else None))
            except Exception:
                pass
        out['recommend_ids'] = ids
    except Exception as e:
        out['recommend_error'] = str(e)

    # Recorre la colección para comprobar si el punto conserva su vector.
    try:
        records, _ = q.scroll(collection_name=QDRANT_COLLECTION, with_vectors=True, with_payload=True, limit=10000)
        out['scrolled_count'] = len(records)
        found = None
        for r in records:
            if getattr(r, 'id', None) == pid or str(getattr(r, 'id', '')) == str(pid):
                found = {'id': getattr(r, 'id', None), 'has_vector': bool(getattr(r, 'vector', None)), 'payload_keys': list((getattr(r,'payload',{}) or {}).keys())}
                break
        out['scrolled_found'] = found
    except Exception as e:
        out['scroll_error'] = str(e)

    # Si hay vector, prueba las firmas históricas de búsqueda y devuelve una muestra.
    try:
        if found and found.get('has_vector'):
            # Recupera el vector exacto del punto encontrado.
            vec = None
            records, _ = q.scroll(collection_name=QDRANT_COLLECTION, with_vectors=True, with_payload=False, limit=10000)
            for r in records:
                if getattr(r, 'id', None) == pid or str(getattr(r, 'id', '')) == str(pid):
                    vec = getattr(r, 'vector', None)
                    break
            if vec is not None:
                search_res = None
                attempts = [
                    lambda q, v, l: q.search(collection_name=QDRANT_COLLECTION, query_vector=v, with_payload=True, limit=l),
                    lambda q, v, l: q.search(collection_name=QDRANT_COLLECTION, vector=v, with_payload=True, limit=l),
                    lambda q, v, l: q.search_points(collection_name=QDRANT_COLLECTION, query_vector=v, with_payload=True, limit=l),
                    lambda q, v, l: q.search_points(collection_name=QDRANT_COLLECTION, vector=v, with_payload=True, limit=l),
                ]
                for fn in attempts:
                    try:
                        r = fn(q, vec, 10)
                        if r is not None:
                            search_res = r
                            break
                    except Exception:
                        continue
                if search_res is not None:
                    # Normaliza las variantes de respuesta del cliente.
                    items = getattr(search_res, 'result', None) or (search_res if isinstance(search_res, list) else None) or (search_res.get('result') if isinstance(search_res, dict) else None)
                    # Como último recurso, materializa una respuesta iterable.
                    if items is None:
                        try:
                            items = list(search_res)
                        except Exception:
                            items = None
                    ids = []
                    if items:
                        for it in items:
                            ids.append(getattr(it, 'id', None) if not isinstance(it, dict) else it.get('id'))
                    out['search_ids'] = ids
                else:
                    out['search_ids'] = []
    except Exception as e:
        out['search_error'] = str(e)

    return out
